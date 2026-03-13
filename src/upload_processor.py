"""
Processamento de uploads para o dashboard do SARE.

Lê arquivos enviados pelo usuário, tenta organizar estruturas comuns
(vindas de Excel, CSV, TXT e Word) e devolve DataFrames padronizados
para alimentar o dashboard.
"""

from __future__ import annotations

import csv
import io
import re
import unicodedata
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

COLUNAS_PADRAO = [
    "id",
    "data",
    "vendedor",
    "setor",
    "produto",
    "categoria",
    "quantidade",
    "valor_total",
    "status",
]

COLUNAS_META_MENSAIS = {
    1: "Meta Jan (R$)",
    2: "Meta Fev (R$)",
    3: "Meta Mar (R$)",
    4: "Meta Abr (R$)",
    5: "Meta Mai (R$)",
    6: "Meta Jun (R$)",
    7: "Meta Jul (R$)",
    8: "Meta Ago (R$)",
    9: "Meta Set (R$)",
    10: "Meta Out (R$)",
    11: "Meta Nov (R$)",
    12: "Meta Dez (R$)",
}

MAPEAMENTO_COLUNAS = {
    "id": "id",
    "codigo": "id",
    "código": "id",
    "data": "data",
    "data venda": "data",
    "data da venda": "data",
    "data_venda": "data",
    "vendedor": "vendedor",
    "nome vendedor": "vendedor",
    "nome do vendedor": "vendedor",
    "vendedor nome": "vendedor",
    "consultor": "vendedor",
    "representante": "vendedor",
    "departamento": "setor",
    "setor": "setor",
    "area": "setor",
    "área": "setor",
    "produto": "produto",
    "serviço": "produto",
    "servico": "produto",
    "item": "produto",
    "solucao": "produto",
    "solução": "produto",
    "categoria": "categoria",
    "grupo": "categoria",
    "tipo": "categoria",
    "quantidade": "quantidade",
    "qtd": "quantidade",
    "qtde": "quantidade",
    "valor": "valor_total",
    "valor total": "valor_total",
    "valor_total": "valor_total",
    "total": "valor_total",
    "status": "status",
    "situação": "status",
    "situacao": "status",
}

MAPEAMENTO_COLUNAS_META = {
    "vendedor": "Vendedor",
    "nome vendedor": "Vendedor",
    "nome do vendedor": "Vendedor",
    "consultor": "Vendedor",
    "representante": "Vendedor",
    "meta": "Meta Mensal (R$)",
    "meta mensal": "Meta Mensal (R$)",
    "meta mensal (r$)": "Meta Mensal (R$)",
    "meta mes": "Meta Mensal (R$)",
    "meta do mes": "Meta Mensal (R$)",
    "meta do mês": "Meta Mensal (R$)",
    "objetivo": "Meta Mensal (R$)",
    "objetivo mensal": "Meta Mensal (R$)",
    "target": "Meta Mensal (R$)",
    "target mensal": "Meta Mensal (R$)",
    "meta jan (r$)": "Meta Jan (R$)",
    "meta fev (r$)": "Meta Fev (R$)",
    "meta mar (r$)": "Meta Mar (R$)",
    "meta abr (r$)": "Meta Abr (R$)",
    "meta mai (r$)": "Meta Mai (R$)",
    "meta jun (r$)": "Meta Jun (R$)",
    "meta jul (r$)": "Meta Jul (R$)",
    "meta ago (r$)": "Meta Ago (R$)",
    "meta set (r$)": "Meta Set (R$)",
    "meta out (r$)": "Meta Out (R$)",
    "meta nov (r$)": "Meta Nov (R$)",
    "meta dez (r$)": "Meta Dez (R$)",
    "meta trimestre (r$)": "Meta Trimestre (R$)",
}

STATUS_MAP = {
    "concluida": "concluida",
    "concluída": "concluida",
    "concluido": "concluida",
    "concluído": "concluida",
    "finalizada": "concluida",
    "finalizado": "concluida",
    "pendente": "pendente",
    "pendencia": "pendente",
    "pendência": "pendente",
    "em aberto": "pendente",
    "cancelada": "cancelada",
    "cancelado": "cancelada",
}


def detectar_extensao_arquivo(arquivo: Any) -> str:
    """
    Detecta a extensão do arquivo enviado.

    A função aceita tanto caminhos locais quanto objetos retornados pelo
    file_uploader do Streamlit, usando o atributo `.name` quando existir.
    """
    nome = getattr(arquivo, "name", str(arquivo))
    return Path(nome).suffix.lower()


def _obter_bytes_arquivo(arquivo: Any) -> bytes:
    """
    Lê o conteúdo binário do arquivo de forma segura.

    Suporta caminhos locais, objetos UploadedFile do Streamlit e buffers em memória.
    """
    if isinstance(arquivo, (str, Path)):
        return Path(arquivo).read_bytes()

    if hasattr(arquivo, "getvalue"):
        return arquivo.getvalue()

    if hasattr(arquivo, "read"):
        conteudo = arquivo.read()
        try:
            arquivo.seek(0)
        except Exception:
            pass
        return conteudo

    raise TypeError("Tipo de arquivo não suportado para leitura.")


def _decodificar_texto(dados: bytes) -> str:
    """
    Decodifica bytes em texto, tentando codificações comuns.

    Isso ajuda a absorver arquivos TXT/CSV salvos em UTF-8, UTF-8 com BOM
    ou Windows-1252, frequentes em ambientes corporativos.
    """
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return dados.decode(encoding)
        except UnicodeDecodeError:
            continue
    return dados.decode("utf-8", errors="ignore")


def _detectar_delimitador(texto: str) -> str:
    """
    Tenta identificar o delimitador mais provável em texto tabular.

    Dá preferência a ';', tab, '|', ',' e cai em ';' como padrão quando
    não há sinal claro de outro separador.
    """
    linhas = [linha for linha in texto.splitlines() if linha.strip()]
    amostra = "\n".join(linhas[:5])
    if not amostra:
        return ";"

    try:
        dialeto = csv.Sniffer().sniff(amostra, delimiters=";,\t|")
        return dialeto.delimiter
    except csv.Error:
        pass

    candidatos = [";", "\t", "|", ","]
    melhor = ";"
    melhor_score = -1
    for candidato in candidatos:
        contagens = [linha.count(candidato) for linha in linhas[:10]]
        score = sum(c > 0 for c in contagens)
        if score > melhor_score:
            melhor_score = score
            melhor = candidato
    return melhor


def _ler_texto_tabular(texto: str) -> pd.DataFrame:
    """
    Converte texto delimitado em DataFrame.

    Aceita conteúdo vindo de CSV/TXT ou de parágrafos extraídos do Word,
    desde que exista um cabeçalho estruturado e delimitadores consistentes.
    """
    delimitador = _detectar_delimitador(texto)
    return pd.read_csv(io.StringIO(texto), sep=delimitador, dtype=str)


def _ler_docx(dados: bytes) -> pd.DataFrame:
    """
    Extrai uma tabela útil de um arquivo DOCX.

    A função tenta primeiro ler a primeira tabela não vazia do documento.
    Se não houver tabela, tenta interpretar os parágrafos como texto delimitado.
    """
    documento = Document(io.BytesIO(dados))

    for tabela in documento.tables:
        linhas = []
        for row in tabela.rows:
            linhas.append([cell.text.strip() for cell in row.cells])
        if len(linhas) >= 2:
            cabecalho = linhas[0]
            corpo = linhas[1:]
            return pd.DataFrame(corpo, columns=cabecalho)

    paragrafos = [p.text.strip() for p in documento.paragraphs if p.text.strip()]
    if paragrafos:
        return _ler_texto_tabular("\n".join(paragrafos))

    return pd.DataFrame()


def ler_arquivo_upload(arquivo: Any) -> pd.DataFrame:
    """
    Lê o arquivo enviado pelo usuário e retorna um DataFrame bruto.

    Formatos suportados nesta versão:
    - .xlsx
    - .csv
    - .txt
    - .docx

    O objetivo aqui é preservar ao máximo a estrutura original para que a
    etapa seguinte faça a organização e padronização dos dados.
    """
    extensao = detectar_extensao_arquivo(arquivo)
    dados = _obter_bytes_arquivo(arquivo)

    if extensao == ".xlsx":
        return pd.read_excel(io.BytesIO(dados), dtype=str)

    if extensao == ".csv":
        return _ler_texto_tabular(_decodificar_texto(dados))

    if extensao == ".txt":
        return _ler_texto_tabular(_decodificar_texto(dados))

    if extensao == ".docx":
        return _ler_docx(dados)

    raise ValueError(
        "Formato não suportado. Use arquivos .xlsx, .csv, .txt ou .docx."
    )


def _normalizar_nome_coluna(coluna: Any) -> str:
    """
    Normaliza nomes de coluna para facilitar mapeamento.

    Remove espaços excedentes, converte para minúsculas e trata variações
    simples de escrita usadas em arquivos manuais.
    """
    nome = str(coluna).strip().lower()
    nome = re.sub(r"\s+", " ", nome)
    return nome


def _splitar_coluna_unica(df: pd.DataFrame) -> pd.DataFrame:
    """
    Tenta expandir planilhas que vieram inteiras em uma única coluna.

    Esse cenário é comum quando o usuário exporta um TXT/CSV e depois salva
    em Excel sem separar corretamente os delimitadores.
    """
    if df.shape[1] != 1:
        return df

    serie = df.iloc[:, 0].dropna().astype(str)
    if serie.empty:
        return df

    amostra = "\n".join(serie.head(10).tolist())
    if not any(sep in amostra for sep in [";", "\t", "|", ","]):
        return df

    linhas = serie.tolist()
    delimitador = _detectar_delimitador(amostra)
    matriz = [linha.split(delimitador) for linha in linhas]
    if len(matriz) < 2:
        return df

    cabecalho = [celula.strip() for celula in matriz[0]]
    corpo = matriz[1:]
    largura = len(cabecalho)
    corpo = [linha[:largura] + [None] * max(0, largura - len(linha)) for linha in corpo]
    return pd.DataFrame(corpo, columns=cabecalho)


def _converter_numero_br(valor: Any) -> float:
    """
    Converte números em formato brasileiro para float.

    Aceita valores como '2.862,10', 'R$ 7.099,04' e também números já nativos.
    """
    if pd.isna(valor):
        return 0.0

    if isinstance(valor, (int, float)):
        return float(valor)

    texto = str(valor).strip()
    if not texto:
        return 0.0

    texto = texto.replace("R$", "").replace(" ", "")
    texto = texto.replace(".", "").replace(",", ".")
    texto = re.sub(r"[^0-9.\-]", "", texto)

    try:
        return float(texto)
    except ValueError:
        return 0.0


def _normalizar_status(status: Any) -> str:
    """
    Padroniza o status das vendas para o padrão interno do SARE.

    O sistema trabalha com três valores principais: 'concluida',
    'pendente' e 'cancelada'. A função também remove acentos para
    absorver variações comuns de escrita.
    """
    texto = str(status).strip().lower()
    texto_sem_acento = unicodedata.normalize("NFKD", texto).encode("ascii", "ignore").decode("ascii")
    return STATUS_MAP.get(texto_sem_acento, STATUS_MAP.get(texto, texto_sem_acento))


def _preparar_base(df: pd.DataFrame) -> pd.DataFrame:
    """
    Aplica limpezas básicas comuns antes da padronização.

    Remove linhas e colunas vazias, além de tentar expandir arquivos que
    vieram concentrados em uma única coluna.
    """
    if df is None or df.empty:
        return pd.DataFrame()

    base = df.copy()
    base = base.dropna(how="all")
    base = base.dropna(axis=1, how="all")
    base = _splitar_coluna_unica(base)
    base.columns = [str(col).strip() for col in base.columns]
    return base


def organizar_planilha_vendas(df: pd.DataFrame) -> pd.DataFrame:
    """
    Organiza e padroniza um DataFrame de vendas para o formato do dashboard.

    A função tenta corrigir cenários comuns de arquivo bagunçado, como:
    - linhas totalmente vazias
    - planilha inteira em uma única coluna
    - nomes de colunas diferentes do padrão interno
    - datas e números em formato textual
    - status escritos com variações

    Se colunas obrigatórias não forem encontradas após a padronização,
    levanta ValueError para que o dashboard avise o usuário.
    """
    if df is None or df.empty:
        return pd.DataFrame(columns=COLUNAS_PADRAO)

    organizado = _preparar_base(df)

    renomeadas = {}
    for coluna in organizado.columns:
        coluna_norm = _normalizar_nome_coluna(coluna)
        if coluna_norm in MAPEAMENTO_COLUNAS:
            renomeadas[coluna] = MAPEAMENTO_COLUNAS[coluna_norm]
    organizado = organizado.rename(columns=renomeadas)

    colunas_obrigatorias = [
        "data",
        "vendedor",
        "produto",
        "categoria",
        "quantidade",
        "valor_total",
        "status",
    ]
    faltando = [col for col in colunas_obrigatorias if col not in organizado.columns]
    if faltando:
        raise ValueError(
            f"Não foi possível identificar as colunas obrigatórias: {faltando}"
        )

    if "setor" not in organizado.columns:
        organizado["setor"] = "Não informado"

    if "id" not in organizado.columns:
        organizado.insert(0, "id", range(1, len(organizado) + 1))

    organizado["data"] = pd.to_datetime(
        organizado["data"],
        errors="coerce",
        dayfirst=True,
    )
    organizado["quantidade"] = pd.to_numeric(
        organizado["quantidade"],
        errors="coerce",
    ).fillna(0).astype(int)
    organizado["valor_total"] = organizado["valor_total"].apply(_converter_numero_br)
    organizado["status"] = organizado["status"].apply(_normalizar_status)

    organizado["vendedor"] = organizado["vendedor"].astype(str).str.strip()
    organizado["setor"] = organizado["setor"].astype(str).str.strip()
    organizado["produto"] = organizado["produto"].astype(str).str.strip()
    organizado["categoria"] = organizado["categoria"].astype(str).str.strip()

    organizado = organizado.dropna(subset=["data", "vendedor", "produto"], how="all")
    organizado = organizado[COLUNAS_PADRAO].copy()
    organizado = organizado.sort_values("data", ascending=False).reset_index(drop=True)
    organizado["id"] = range(1, len(organizado) + 1)

    return organizado


def organizar_planilha_metas(
    df: pd.DataFrame,
    data_referencia: pd.Timestamp | None = None,
) -> pd.DataFrame:
    """
    Organiza e padroniza um DataFrame de metas para o formato esperado pelo SARE.

    A função aceita arquivos com coluna de meta genérica (como 'Meta' ou
    'Meta Mensal') e também arquivos com metas por mês já separadas.

    Se não for possível identificar pelo menos o vendedor e uma coluna de
    meta, levanta ValueError para que o dashboard decida o melhor fallback.
    """
    if df is None or df.empty:
        return pd.DataFrame(columns=["Vendedor", "Meta Mensal (R$)"])

    base = _preparar_base(df)

    renomeadas = {}
    for coluna in base.columns:
        coluna_norm = _normalizar_nome_coluna(coluna)
        if coluna_norm in MAPEAMENTO_COLUNAS_META:
            renomeadas[coluna] = MAPEAMENTO_COLUNAS_META[coluna_norm]
    base = base.rename(columns=renomeadas)

    if "Vendedor" not in base.columns:
        raise ValueError("Não foi possível identificar a coluna de vendedor nas metas.")

    colunas_meta = [
        col for col in base.columns
        if col == "Meta Mensal (R$)"
        or col in COLUNAS_META_MENSAIS.values()
        or col == "Meta Trimestre (R$)"
    ]
    if not colunas_meta:
        raise ValueError("Não foi possível identificar nenhuma coluna de meta.")

    metas = base[["Vendedor", *colunas_meta]].copy()
    metas["Vendedor"] = metas["Vendedor"].astype(str).str.strip()
    metas = metas[metas["Vendedor"] != ""]

    for coluna in colunas_meta:
        metas[coluna] = metas[coluna].apply(_converter_numero_br)

    if data_referencia is not None and "Meta Mensal (R$)" not in metas.columns:
        coluna_mes = COLUNAS_META_MENSAIS.get(data_referencia.month)
        if coluna_mes and coluna_mes in metas.columns:
            metas.attrs["meta_coluna"] = coluna_mes

    metas.attrs["origem_metas"] = "upload"
    return metas.reset_index(drop=True)


def construir_metas_demonstrativas(
    df_vendas: pd.DataFrame,
    multiplicador: float = 1.15,
    piso_meta: float = 5000.0,
) -> pd.DataFrame:
    """
    Constrói metas demonstrativas quando o upload não traz metas explícitas.

    O objetivo não é substituir metas reais de negócio, e sim evitar que o
    dashboard misture vendedores do upload com a planilha padrão antiga.
    As metas são derivadas do realizado mais recente por vendedor e ficam
    marcadas nos atributos do DataFrame como origem demonstrativa.
    """
    if df_vendas is None or df_vendas.empty:
        return pd.DataFrame(columns=["Vendedor", "Meta Mensal (R$)"])

    base = df_vendas.copy()
    if "data" in base.columns:
        base["data"] = pd.to_datetime(base["data"], errors="coerce")
        datas_validas = base["data"].dropna()
        if not datas_validas.empty:
            data_referencia = datas_validas.max()
            base = base[
                (base["data"].dt.year == data_referencia.year)
                & (base["data"].dt.month == data_referencia.month)
            ]
    if "status" in base.columns:
        base = base[base["status"] == "concluida"]

    if "vendedor" not in base.columns or "valor_total" not in base.columns:
        return pd.DataFrame(columns=["Vendedor", "Meta Mensal (R$)"])

    ranking = (
        base.groupby("vendedor", as_index=False)["valor_total"]
        .sum()
        .rename(columns={"vendedor": "Vendedor"})
    )
    if ranking.empty:
        vendedores = pd.DataFrame(
            {"Vendedor": sorted(df_vendas["vendedor"].dropna().astype(str).str.strip().unique())}
        )
        vendedores["Meta Mensal (R$)"] = piso_meta
        vendedores.attrs["origem_metas"] = "demonstrativa"
        vendedores.attrs["descricao_origem"] = "Metas demonstrativas geradas automaticamente a partir do upload."
        return vendedores

    ranking["Meta Mensal (R$)"] = (
        ranking["valor_total"].fillna(0.0).astype(float) * multiplicador
    ).clip(lower=piso_meta).round(2)
    metas = ranking[["Vendedor", "Meta Mensal (R$)"]].copy()
    metas.attrs["origem_metas"] = "demonstrativa"
    metas.attrs["descricao_origem"] = "Metas demonstrativas geradas automaticamente a partir do upload."
    return metas.sort_values("Vendedor").reset_index(drop=True)


def preparar_upload_vendas(arquivo: Any) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    """
    Lê o arquivo enviado e devolve versões bruta e organizada.

    O retorno inclui um dicionário de metadados simples para que o dashboard
    possa informar origem, formato e quantidade de registros processados.
    """
    bruto = ler_arquivo_upload(arquivo)
    organizado = organizar_planilha_vendas(bruto)
    info = {
        "nome_arquivo": getattr(arquivo, "name", str(arquivo)),
        "extensao": detectar_extensao_arquivo(arquivo),
        "linhas_brutas": int(len(bruto)),
        "linhas_organizadas": int(len(organizado)),
        "colunas_brutas": bruto.columns.tolist(),
    }
    return bruto, organizado, info


def preparar_upload_metas(
    arquivo: Any,
    data_referencia: pd.Timestamp | None = None,
) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    """
    Lê o arquivo enviado e devolve versões bruta e organizada de metas.

    Essa função é usada quando o usuário envia um arquivo de metas separado.
    """
    bruto = ler_arquivo_upload(arquivo)
    organizado = organizar_planilha_metas(bruto, data_referencia=data_referencia)
    info = {
        "nome_arquivo": getattr(arquivo, "name", str(arquivo)),
        "extensao": detectar_extensao_arquivo(arquivo),
        "linhas_brutas": int(len(bruto)),
        "linhas_organizadas": int(len(organizado)),
        "colunas_brutas": bruto.columns.tolist(),
    }
    return bruto, organizado, info


def dataframe_para_excel_bytes(df: pd.DataFrame, sheet_name: str = "Vendas_Organizadas") -> bytes:
    """
    Converte um DataFrame em arquivo Excel em memória.

    Isso permite ao dashboard oferecer download da versão organizada sem
    precisar gravar arquivos temporários em disco.
    """
    buffer = io.BytesIO()
    exportar = df.copy()
    if "data" in exportar.columns and pd.api.types.is_datetime64_any_dtype(exportar["data"]):
        exportar["data"] = exportar["data"].dt.strftime("%d/%m/%Y")

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        exportar.to_excel(writer, index=False, sheet_name=sheet_name)
    return buffer.getvalue()
